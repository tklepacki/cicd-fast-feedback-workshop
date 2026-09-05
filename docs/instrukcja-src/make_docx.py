# -*- coding: utf-8 -*-
"""Buduje .docx z instrukcją przygotowania, w oprawie zbliżonej do poprzedniej edycji."""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x33, 0x5E)
GREY = RGBColor(0x44, 0x44, 0x4C)

def shade(cell_or_par, hexcolor):
    el = cell_or_par._tc if hasattr(cell_or_par, '_tc') else cell_or_par._p.get_or_add_pPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor)
    el.append(sh)

doc = Document()
for s in doc.styles['Normal'], :
    s.font.name = 'Calibri'; s.font.size = Pt(10.5)
for section in doc.sections:
    section.top_margin = section.bottom_margin = Inches(0.7)
    section.left_margin = section.right_margin = Inches(0.8)

# nagłówek
bar = doc.add_paragraph(); shade(bar, '1F335E')
r = bar.add_run('SZYBKI FEEDBACK W CI/CD')
r.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.size = Pt(11)

t = doc.add_paragraph(); rt = t.add_run('Instrukcja przygotowania do warsztatów')
rt.bold = True; rt.font.size = Pt(22); rt.font.color.rgb = NAVY

sub = doc.add_paragraph()
rs = sub.add_run('Pipeline, który testuje i raportuje bez utraty jakości  ·  GitHub Actions')
rs.font.size = Pt(11); rs.font.color.rgb = GREY

def code(text):
    p = doc.add_paragraph(); shade(p, 'F2F3F5')
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.font.name = 'Consolas'; r.font.size = Pt(9.5)
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')

def note(text, fill='FFF6E0'):
    p = doc.add_paragraph(); shade(p, fill)
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    for bold, chunk in parse_bold(text):
        r = p.add_run(chunk); r.font.size = Pt(10); r.bold = bold

def parse_bold(text):
    out = []
    for part in re.split(r'(\*\*[^*]+\*\*)', text):
        if not part: continue
        if part.startswith('**') and part.endswith('**'):
            out.append((True, part[2:-2]))
        else:
            out.append((False, part.replace('`', '')))
    return out

def para(text, bullet=False, size=10.5):
    p = doc.add_paragraph(style='List Bullet' if bullet else None)
    p.paragraph_format.space_after = Pt(4)
    for bold, chunk in parse_bold(text):
        r = p.add_run(chunk); r.font.size = Pt(size); r.bold = bold
    return p

def heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True; r.font.color.rgb = NAVY
    r.font.size = Pt(15 if level == 1 else 12)
    return p

def table(rows, header=True):
    tb = doc.add_table(rows=0, cols=len(rows[0]))
    tb.style = 'Table Grid'; tb.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, row in enumerate(rows):
        cells = tb.add_row().cells
        for c, val in zip(cells, row):
            c.text = ''
            p = c.paragraphs[0]
            for bold, chunk in parse_bold(val):
                r = p.add_run(chunk); r.font.size = Pt(10); r.bold = bold or (header and i == 0)
            if header and i == 0:
                shade(c, 'E8EBF0')
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ---------------- treść ----------------
para("Warsztat jest praktyczny — przez większość dnia będziesz pracować przy klawiaturze "
     "na własnym repozytorium. Żeby nie stracić pierwszej godziny na instalacje, "
     "**wykonaj poniższe kroki przed warsztatem** i sprawdź, że wszystko działa. "
     "Całość zajmuje około 30 minut.")
note("**Przyjdź z laptopem**, na którym masz uprawnienia do instalowania oprogramowania "
     "i dostęp do sieci bez blokad — firmowy VPN albo polityki bezpieczeństwa potrafią "
     "zablokować npm lub GitHuba.")

heading("1. Narzędzia do zainstalowania")
table([["Narzędzie", "Wersja", "Skąd"],
       ["Visual Studio Code", "dowolna aktualna", "code.visualstudio.com/download"],
       ["Node.js", "**22 LTS lub nowszy**", "nodejs.org"],
       ["Git", "dowolna aktualna", "git-scm.com/downloads"]])
para("Sprawdź w terminalu, czy wszystko odpowiada:")
code("node --version     # v22.x lub nowsza\nnpm --version\ngit --version")
note("**Nie potrzebujesz Javy ani Dockera.** Jeśli brałeś udział w poprzedniej edycji tego "
     "warsztatu — tym razem pracujemy wyłącznie na GitHub Actions.", 'E9F2E9')

heading("2. Konto GitHub")
para("Potrzebujesz **prywatnego konta GitHub**. Darmowe w zupełności wystarczy.")
para("Jeśli masz konto firmowe objęte logowaniem SSO, **użyj konta prywatnego** — firmowe "
     "polityki potrafią blokować tworzenie repozytoriów publicznych i uruchamianie workflow.")

heading("3. Utwórz własną kopię repozytorium")
para("Otwórz repozytorium warsztatowe:")
code("https://github.com/tklepacki/cicd-fast-feedback-workshop")
para("Kliknij zielony przycisk **Use this template** → **Create a new repository**, "
     "a w formularzu ustaw:")
table([["Pole", "Ustawienie"],
       ["Repository name", "dowolna, np. warsztat-cicd"],
       ["Widoczność", "**Public**"],
       ["**Include all branches**", "**zaznacz**"]])
heading("Dwie rzeczy, które łatwo przeoczyć", 2)
note("**Zaznacz „Include all branches”.** Bez tego nie dostaniesz branchy solution/* "
     "z rozwiązaniami ani branchy demo/* z celowo zepsutym kodem — a używamy ich już "
     "w pierwszym zadaniu. To najczęściej pomijany krok w całej instrukcji.", 'FDECEA')
note("**Repozytorium musi być publiczne.** Nie chodzi o dzielenie się kodem, tylko o to, "
     "że na darmowym koncie repozytorium publiczne dostaje nielimitowane minuty GitHub Actions "
     "i mocniejszy runner — 4 rdzenie zamiast 2. Na repozytorium prywatnym część zadań "
     "jest technicznie niewykonalna.", 'FDECEA')

heading("4. Sklonuj repozytorium i zainstaluj zależności")
code("git clone <adres-Twojego-nowego-repozytorium>\ncd <nazwa-katalogu>\n"
     "npm ci\nnpx playwright install --with-deps chromium")
note("Instalacja przeglądarki jest tym razem **potrzebna** — w warsztacie są testy UI, "
     "nie tylko API. Pobranie Chromium zajmuje chwilę, więc zrób to przed warsztatem.")

heading("5. Sprawdź, że wszystko działa")
code("npm run verify")
para("Uruchamia lint, typecheck, build i testy jednostkowe. Powinno zakończyć się bez błędów "
     "i wypisać 78 passed.")
para("Sprawdź też testy przeglądarkowe i samą aplikację:")
code("npm run test:smoke     # 5 testów, powinny przejść\nnpm run dev            # http://localhost:5173")
para("Aplikacja to prosty sklep: katalog, koszyk, zamówienie. Dokumentacja API jest pod "
     "adresem localhost:3000/api/docs po uruchomieniu npm start.")

heading("6. Włącz GitHub Actions w swoim repozytorium")
para("Wejdź w swoje repozytorium na GitHubie → zakładka **Actions**. Jeśli zobaczysz komunikat "
     "o wyłączonych workflow, kliknij przycisk włączający je. Bez tego pierwsze zadanie nie ruszy.")
para("Następnie sprawdź, że pipeline działa:")
para("wejdź w Actions → wybierz workflow **CI**", True)
para("jeśli nie ma żadnego przebiegu, zrób dowolny commit i wypchnij go na main", True)
para("poczekaj — przebieg powinien być zielony i trwać około czterech minut", True)
note("**Te cztery minuty to nie błąd.** To jest punkt wyjścia, który będziemy skracać "
     "przez cały dzień.", 'E9F2E9')

heading("Lista kontrolna")
for item in [
    "node --version pokazuje 22 lub nowszą",
    "git --version działa",
    "Visual Studio Code zainstalowany",
    "mam prywatne konto GitHub",
    "utworzyłem repozytorium z szablonu, **publiczne**, z **Include all branches**",
    "widzę u siebie branche solution/zadanie-01 i demo/failing-lint",
    "npm ci przeszło bez błędów",
    "npx playwright install --with-deps chromium przeszło bez błędów",
    "npm run verify kończy się sukcesem (78 testów)",
    "npm run test:smoke kończy się sukcesem (5 testów)",
    "Actions są włączone i pierwszy przebieg zakończył się na zielono",
]:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run('☐  '); r.font.size = Pt(11)
    for bold, chunk in parse_bold(item):
        rr = p.add_run(chunk); rr.font.size = Pt(10.5); rr.bold = bold

heading("Gdyby coś nie zadziałało")
table([["Objaw", "Co zrobić"],
       ["npm ci zgłasza błąd wersji Node", "sprawdź node --version; projekt wymaga 22 lub nowszej"],
       ["npx playwright install przerywa się", "najczęściej blokada sieci firmowej — spróbuj z innej sieci lub bez VPN"],
       ["nie widzę branchy solution/* ani demo/*", "kopia powstała bez „Include all branches” — usuń repozytorium i utwórz je ponownie"],
       ["zakładka Actions pusta, workflow nie startuje", "sprawdź, czy Actions są włączone (punkt 6) i czy repozytorium jest publiczne"]])
para("Jeśli utkniesz — napisz do mnie przed warsztatem, żebyśmy nie tracili na to czasu na sali.")

out = '/Users/tomaszklepacki/Desktop/workshop/docs/Instrukcja przygotowania - Szybki feedback w CICD.docx'
doc.save(out)
print('zapisano:', out)
